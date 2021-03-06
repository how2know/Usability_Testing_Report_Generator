from docx.document import Document
from bs4 import BeautifulSoup
from typing import List, Dict, Union
import numpy as np
import pandas as pd
from docx.shared import Cm

from docx_package.dropdown_lists import DropDownLists
from docx_package.results import ResultsChapter
from docx_package.picture import Picture
from eye_tracking_package.plot import Plot


class TimeOnTasks:
    """
    Class that represents the 'Time on tasks' chapter and the visualization of its results.
    """

    # name of table as it appears in the tables list
    TIME_ON_TASK_TABLE = 'Time on tasks table'
    PLOT_TYPE_TABLE = 'Time on tasks plot type table'
    DECISION_TABLE = 'Time on tasks decision table'

    # parameter keys as they appear in the parameters dictionary
    PARTICIPANTS_NUMBER_KEY = 'Number of participants'
    TASKS_NUMBER_KEY = 'Number of critical tasks'

    # column labels as they appear in the Tobii data frame
    EVENT_LABEL = 'Event'
    SECONDS_LABEL = 'Seconds'

    # information about the headings of this chapter
    TITLE = 'Time on tasks'
    TITLE_STYLE = 'Heading 2'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # path to plot image files
    PARTICIPANT_FIGURE_PATH = 'Outputs/Time_on_task_participant{}.png'
    BAR_PLOT_FIGURE_PATH = 'Outputs/Time_on_task_bar_plot.png'
    BOX_PLOT_FIGURE_PATH = 'Outputs/Time_on_task_box_plot.png'

    # caption of the pie plot figure
    BAR_PLOT_CAPTION = 'Bar plot showing the mean of the task completion times and the 95% confidence interval.'
    BOX_PLOT_CAPTION = 'Box plot showing the mean, the 25% and 75% quartiles, and the distribution of the task completion times.'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]],
                 tobii_data: pd.DataFrame
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value).
            tobii_data: Data frame that contains the given Tobii data.
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.picture_paths = picture_paths_list
        self.parameters = parameters_dictionary
        self.tables = list_of_tables
        input_table_index = self.tables.index(self.TIME_ON_TASK_TABLE)
        self.input_table = text_input_document.tables[input_table_index]
        self.tobii_data = tobii_data

    @ property
    def tasks_number(self) -> int:
        """
        Returns:
            Biggest number of critical tasks between the one given in the 'Number of critical tasks' input,
            the one that corresponds to the time on tasks input table,
            and the one given in the Tobii data.
        """

        table_tasks_number = 0
        tobii_tasks_number = 0

        # get the index of the last row that is filled which corresponds to the number of critical tasks
        for i in range(1, len(self.input_table.rows)):
            row_filled = False
            for cell in self.input_table.rows[i].cells[1:]:
                if cell.text:
                    row_filled = True
            if row_filled:
                table_tasks_number = i

        # get the number of critical tasks described in the Tobii data
        tobii_tasks = self.tobii_data[self.EVENT_LABEL].tolist()
        for task in tobii_tasks:
            number = int(task.replace('Task', ''))
            if number > tobii_tasks_number:
                tobii_tasks_number = number

        return max([tobii_tasks_number,
                    table_tasks_number,
                    self.parameters[self.TASKS_NUMBER_KEY]])

    @ property
    def participants_number(self) -> int:
        """
        Returns:
            Biggest number of participants between the one given in the 'Number of participants' input,
            the one that corresponds to the time on tasks input table,
            and the one given in the Tobii data.
        """

        table_participants_number = 0
        tobii_participants_number = 0

        # get the index of the last column that is filled which corresponds to the number of participants
        for j in range(1, len(self.input_table.columns)):
            column_completed = False
            for cell in self.input_table.columns[j].cells[1:]:
                if cell.text:
                    column_completed = True
            if column_completed:
                table_participants_number = j

        # get the number of participants described in the Tobii data
        tobii_participants = self.tobii_data.index.values.tolist()
        for participant in tobii_participants:
            number = int(participant.replace('Participant', ''))
            if number > tobii_participants_number:
                tobii_participants_number = number

        return max([tobii_participants_number,
                    table_participants_number,
                    self.parameters[self.PARTICIPANTS_NUMBER_KEY]])

    @ property
    def tasks(self) -> List[str]:
        """
        Returns:
            List of task names.
        """

        tasks = []
        for i in range(1, self.tasks_number + 1):

            # write the real task name if it is provided or a generic name if not
            try:
                tasks.append(self.parameters['Critical task {} name'.format(i)])
            except KeyError:
                tasks.append('Critical task {}'.format(i))

        return tasks

    @ property
    def participants(self) -> List[str]:
        """
        Returns:
            List of participants, i.e. [Participant 1, Participant 2, ...].
        """

        participants = ['Participant{}'.format(i) for i in range(1, self.participants_number + 1)]
        return participants

    @ property
    def times_from_table(self) -> pd.DataFrame:
        """
        Returns:
            Data frame of tasks completion times with participants as index and task names as columns.
        """

        rows_number = self.tasks_number
        columns_number = self.participants_number

        # create a matrix full of NaN of the size of the input table
        times_matrix = np.full((rows_number, columns_number), np.nan)
        for i in range(rows_number):
            for j in range(columns_number):

                # set the completion time or keep NaN if no time were given
                try:
                    time = float(self.input_table.cell(i+1, j+1).text)
                    times_matrix[i, j] = time
                except ValueError:
                    pass

        # create a data frame with the transposed matrix to have participants as rows and tasks as columns
        times_df = pd.DataFrame(times_matrix.transpose(), index=self.participants, columns=self.tasks)

        return times_df

    def times_from_tables_and_tobii(self):
        """
        Complete the data frame of tasks completion times with the input given through Tobii.

        Returns:
            Data frame of tasks completion times with participants as index and task names as columns.
        """

        # data frame of tasks completion times with the input given through the text input form
        times_df = self.times_from_table

        for participant in self.participants:
            participants_data = self.tobii_data[self.tobii_data.index == participant]
            for idx, task in enumerate(self.tasks):
                tasks_data = participants_data[participants_data[self.EVENT_LABEL] == 'Task{}'.format(idx+1)]

                # replace the time in the data frame or do nothing when no time was given through Tobii
                try:
                    seconds = tasks_data[self.SECONDS_LABEL].to_numpy()
                    time = seconds[1] - seconds[0]
                    times_df.loc[participant].at[task] = time
                except IndexError:
                    pass

        # delete all rows and columns that are full of missing values
        times_df = times_df.dropna(axis=0, how='all')
        times_df = times_df.dropna(axis=1, how='all')

        return times_df

    @ property
    def plot_type(self) -> str:
        """
        Returns:
            Dropdown list value of the parameter table corresponding to the plot type,
            i.e. 'Bar plot' or 'Box plot'.
        """

        plot_type_list = DropDownLists.get_from_table(self.text_input_soup,
                                                       self.tables.index(self.PLOT_TYPE_TABLE)
                                                       )
        return plot_type_list[0]

    def make_plots(self):
        """
        Create plots with the time on tasks values.

        One bar plot for each participant is created.
        One bar plot showing a confidence interval of 95% and one box plot with the data of
        all participants are created.
        """

        task_times_df = self.times_from_tables_and_tobii()

        # create a bar plot for each participant
        for idx, participant in enumerate(self.participants):

            try:
                participant_times = task_times_df.loc[participant].to_numpy()
                participant_times_df = pd.DataFrame(data=[participant_times],
                                                    columns=task_times_df.columns)

                Plot.make_barplot(data_frame=participant_times_df,
                                  figure_save_path=self.PARTICIPANT_FIGURE_PATH.format(idx+1),
                                  title='Time on task: participant {}'.format(idx+1),
                                  ylabel='Completion time [s]',
                                  xlabel='Critical task'
                                  )

            # pass when no data is available for a participant
            except KeyError:
                pass

        # create a bar plot and a box plot with the data of all participants or do nothing if no data is provided
        try:
            Plot.make_barplot(data_frame=task_times_df,
                              figure_save_path=self.BAR_PLOT_FIGURE_PATH,
                              title='Time on task',
                              ylabel='Completion time [s]',
                              xlabel='Critical task'
                              )
            Plot.make_boxplot(data_frame=task_times_df,
                              figure_save_path=self.BOX_PLOT_FIGURE_PATH,
                              title='Time on task',
                              ylabel='Completion time [s]',
                              xlabel='Critical task'
                              )
        except ValueError:
            pass

    def write_chapter(self):
        """
        Write the whole chapter 'Time on tasks', including the chosen plot.
        """

        decision_table_index = self.tables.index(self.DECISION_TABLE)
        decision = DropDownLists.get_from_table(self.text_input_soup, decision_table_index)

        if decision[0] == 'Yes':
            self.make_plots()

            time_on_tasks = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                           self.tables, self.picture_paths, self.parameters)

            self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)

            # add bar plot or box plot depending on the choice of plot type or do nothing if no data is provided
            try:
                if self.plot_type == 'Bar plot':
                    Picture.add_picture_and_caption(self.report,
                                                    [self.BAR_PLOT_FIGURE_PATH],
                                                    self.BAR_PLOT_FIGURE_PATH,
                                                    self.BAR_PLOT_CAPTION,
                                                    width=Cm(12)
                                                    )
                if self.plot_type == 'Box plot':
                    Picture.add_picture_and_caption(self.report,
                                                    [self.BOX_PLOT_FIGURE_PATH],
                                                    self.BOX_PLOT_FIGURE_PATH,
                                                    self.BOX_PLOT_CAPTION,
                                                    width=Cm(12)
                                                    )
            except FileNotFoundError:
                pass

            self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
            time_on_tasks.write_chapter()
