from docx.document import Document
from docx.table import Table
from docx.shared import Cm
from typing import List, Dict, Union
from bs4 import BeautifulSoup

from docx_package.dropdown_lists import DropDownLists
from docx_package.picture import Picture


class Chapter:
    """
    Class that represents all classical chapter.

    Every classical chapter are represented by this class, i.e. all chapters except
    'Terms definitions' and those of the 'Results' section.
    """

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 title: str,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            title: Title of the chapter.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value).
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.title = title
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters_dictionary = parameters_dictionary

    def heading_index(self) -> int:
        """
        Returns:
            Paragraph index of the chapter heading in the text input document.
        """

        for paragraph_index, paragraph in enumerate(self.text_input.paragraphs):
            if paragraph.text == self.title and 'Heading' in paragraph.style.name:
                return paragraph_index

    @ property
    def paragraph_table(self) -> Table:
        """
        Returns:
            Table of the input .docx file where the paragraphs of the chapter are written.
        """

        problem_table_index = self.tables.index('{} text table'.format(self.title))
        return self.text_input.tables[problem_table_index]

    def paragraphs_from_table(self) -> List[str]:
        """
        Returns:
            List of all paragraphs (as text string) of the chapter.
        """

        cell = self.paragraph_table.rows[1].cells[0]

        list_of_paragraphs = []

        for paragraph in cell.paragraphs:
            list_of_paragraphs.append(paragraph.text)

        return list_of_paragraphs

    @ property
    def chapter_parameters(self) -> List[str]:
        """
        Returns:
            List of parameters needed to be writen in the chapter.
        """

        return DropDownLists.get_from_table(self.text_input_soup,
                                            self.tables.index('{} parameter table'.format(self.title))
                                            )

    @ property
    def picture_name(self) -> str:
        """
        Returns:
            Title with underscores instead of spaces, e.g. 'Use environment' becomes 'Use_environment'.
        """

        return self.title.replace(' ', '_')

    @ property
    def picture_captions(self) -> List[str]:
        """
        Returns:
            List of the captions of the pictures.
        """

        # read the caption text from the corresponding table in text input and append it to a list
        captions_list = []
        table_index = self.tables.index('{} caption table'.format(self.title))
        table = self.text_input.tables[table_index]
        for i in range(1, 4):
            cell = table.cell(i, 1)
            captions_list.append(cell.text)

        return captions_list

    def add_picture(self):
        """
        Add max. 3 pictures to the chapter.
        """

        captions = self.picture_captions
        picture_name = self.picture_name

        # add pictures that correspond to the given picture file names with the corresponding captions
        for i in range(0, 3):
            Picture.add_picture_and_caption(self.report,
                                            self.picture_paths,
                                            picture_name + str(i+1),
                                            captions[i],
                                            width=Cm(10)
                                            )

    def write_chapter(self):
        """
        Write the whole chapter.

        Write the heading, the paragraphs including the parameters, the pictures and their caption.
        """

        # write heading with the corresponding style
        heading_style = self.text_input.paragraphs[self.heading_index()].style.name
        self.report.add_paragraph(self.title, heading_style)

        # stores values of corresponding parameter keys in a list
        parameters_values = ['', '', '']
        for parameter_idx, parameter in enumerate(self.chapter_parameters):
            if parameter != '-':
                parameters_values[parameter_idx] = self.parameters_dictionary[parameter]

        # write paragraphs including values of parameters or print an error message when a wrong reference was given
        for paragraph in self.paragraphs_from_table():
            try:
                self.report.add_paragraph(
                    paragraph.format(parameters_values[0], parameters_values[1], parameters_values[2]),
                    'Normal'
                )
            except (IndexError, ValueError) as error:
                self.report.add_paragraph(paragraph, 'Normal')
                print('A wrong reference to a parameter was given in the chapter "{}"'.format(self.title), ': ', error)

        # add pictures and their caption at the end of the chapter
        self.add_picture()
