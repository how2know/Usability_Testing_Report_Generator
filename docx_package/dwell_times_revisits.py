from docx.document import Document
from docx.table import Table
from bs4 import BeautifulSoup
from typing import List, Dict, Union
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
import pandas as pd

from docx_package import layout
from docx_package.results import ResultsChapter
from txt_package import plot, eye_tracking


class DwellTimesAndRevisits:
    """
    Class that represents the 'Dwells times and revisits' chapter and the visualization of its results.
    """

    # information about the headings of this chapter
    TITLE = 'Dwell times and revisits'
    TITLE_STYLE = 'Heading 2'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # index name of the statistics
    SUM_INDEX = 'Sum'
    MEAN_INDEX = 'Mean'
    MAX_INDEX = 'Max'
    MIN_INDEX = 'Min'

    # text of the cells of the first row
    TABLE_FIRST_ROW = ['AOI', 'Dwell times [s]', 'Average [s]', 'Max [s]', 'Min [s]', 'Revisits']

    # color for cell shading
    LIGHT_GREY_10 = 'D0CECE'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]],
                 list_of_dataframes: List[pd.DataFrame]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
            list_of_dataframes: List of data frames containing the cGOM data of each participant
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables
        self.parameters = parameters_dictionary
        self.cGOM_dataframes = list_of_dataframes

    # TODO: change docstrings and function name (because it creates plot)
    def dwell_times_stat(self) -> pd.DataFrame:
        """
        Returns:
            Data frame containing the dwell times of all participants (index) and all AOI (columns).
        """

        # main dwell times data frame that contains the dwell times of all participants
        all_dwell_times_df = pd.DataFrame()

        # create a data frame with the dwell times and statistics of each participants
        # and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = eye_tracking.areas_of_interest(dataframe)
            participants_df = eye_tracking.dwell_times(aois, dataframe)
            all_dwell_times_df = all_dwell_times_df.append(participants_df)

            # TODO: use function for plotting
            # plot the total sum of the dwell times for each participants
            sums = participants_df[self.SUM_INDEX].to_numpy()
            sums = sums[~np.isnan(sums)]
            plt.pie(sums, labels=aois)
            plt.show()

        # create a data frame with the mean of the statistics for all participants for each AOI
        all_aois = eye_tracking.areas_of_interest(all_dwell_times_df)
        dwell_times_table = pd.DataFrame(index=all_aois,
                                         columns=[self.SUM_INDEX, self.MEAN_INDEX, self.MAX_INDEX, self.MIN_INDEX],
                                         data=np.zeros((len(all_aois), 4))
                                         )

        # calculate the mean of all statistics and store it in the data frame
        for aoi in all_aois:
            data_of_aoi = all_dwell_times_df[all_dwell_times_df.index == aoi]
            dwell_times_table.loc[aoi].at[self.SUM_INDEX] = data_of_aoi[self.SUM_INDEX].mean()
            dwell_times_table.loc[aoi].at[self.MEAN_INDEX] = data_of_aoi[self.MEAN_INDEX].mean()
            dwell_times_table.loc[aoi].at[self.MAX_INDEX] = data_of_aoi[self.MAX_INDEX].mean()
            dwell_times_table.loc[aoi].at[self.MIN_INDEX] = data_of_aoi[self.MIN_INDEX].mean()

        print(dwell_times_table)

        # TODO: create main plot

        return dwell_times_table

    def revisits_stat(self) -> pd.DataFrame:
        """
        Returns:
            Data frame containing the revisits of all participants (index) and all tasks (columns).
            The last row contains the mean revisits for each task.
        """

        # main revisits data frame
        revisits_df = pd.DataFrame()

        # create a data frame with the revisits for each participant and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = eye_tracking.areas_of_interest(dataframe)
            revisits = eye_tracking.revisits(aois, dataframe)
            participant_revisits = pd.DataFrame(index=['Participant {}'.format(idx + 1)],
                                                columns=aois,
                                                data=[revisits]
                                                )
            revisits_df = revisits_df.append(participant_revisits)

        # calculate the mean of revisits for each AOI and append it to the main data frame
        revisits_mean = revisits_df.mean().to_numpy()
        revisits_mean_df = pd.DataFrame(index=['Mean'],
                                        columns=revisits_df.columns,
                                        data=[revisits_mean]
                                        )
        revisits_df = revisits_df.append(revisits_mean_df)

        return revisits_df

    # TODO: write better and comment
    def add_table(self):
        """
        Create the table with
        """

        # data frames containing the data
        dwell_times_df = self.dwell_times_stat()
        revisits_df = self.revisits_stat()

        # areas of interests
        aois = dwell_times_df.index

        # create table
        table = self.report.add_table(len(aois) + 1, len(self.TABLE_FIRST_ROW))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # write the first row
        for index, label in enumerate(self.TABLE_FIRST_ROW):
            cell = table.rows[0].cells[index]
            cell.text = label
            layout.set_cell_shading(cell, self.LIGHT_GREY_10)  # color the cell in light_grey_10
            cell.paragraphs[0].runs[0].font.bold = True

        # write the first column with the name of the areas of interest
        for idx, aoi in enumerate(aois):
            cell = table.columns[0].cells[idx+1]
            cell.text = aoi

        matrix = dwell_times_df.to_numpy()
        for i in range(len(aois)):
            for j in range(4):
                table.cell(i+1, j+1).text = str(round(matrix[i, j], 4))

        # write the third column with the mean of revisits of each AOI
        for idx, revisits in enumerate(revisits_df.loc['Mean'].to_numpy()):
            cell = table.columns[5].cells[idx+1]
            cell.text = str(round(revisits, 4))

        # set the vertical and horizontal alignment of all cells
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # TODO: do we need these lines
                # cell.paragraphs[0].style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # cell.paragraphs[0].style.name = 'Table'

    def write_chapter(self):
        """
        Write the whole chapter 'Dwell times and revisits', including the table.
        """

        time_on_tasks = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                       self.tables, self.parameters)

        self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)

        self.add_table()

        self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
        time_on_tasks.write_chapter()

    '''
    def dwell_times_stat1(self) -> pd.DataFrame:
        """
        Returns:
            Data frame containing the dwell times of all participants (index) and all AOI (columns).
            The last row contains the mean dwell times for each AOI.
        """

        # main dwell times data frame
        dwell_times_df = pd.DataFrame()

        # create a data frame with the dwell times for each participant and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = eye_tracking.areas_of_interest(dataframe)
            dwell_times = eye_tracking.dwell_times1(aois, dataframe)
            participant_dwell_times = pd.DataFrame(index=['Participant {}'.format(idx+1)],
                                                   columns=aois,
                                                   data=dwell_times.reshape(1, -1)
                                                   )
            dwell_times_df = dwell_times_df.append(participant_dwell_times)

        # calculate the mean of dwell times for each AOI and append it to the main data frame
        dwell_times_mean = dwell_times_df.mean().to_numpy()
        dwell_times_mean_df = pd.DataFrame(index=['Mean'],
                                           columns=dwell_times_df.columns,
                                           data=[dwell_times_mean]
                                           )
        dwell_times_df = dwell_times_df.append(dwell_times_mean_df)

        return dwell_times_df
    '''

    '''
        def add_table1(self):
            """
            Create the table with
            """

            # data frames containing the data
            dwell_times_df = self.dwell_times_stat1()
            revisits_df = self.revisits_stat()

            # areas of interests
            aois = dwell_times_df.columns

            # create table
            table = self.report.add_table(len(aois) + 1, 3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # write the first row
            for index, label in enumerate(self.TABLE_FIRST_ROW2):
                cell = table.rows[0].cells[index]
                cell.text = label
                layout.set_cell_shading(cell, self.LIGHT_GREY_10)  # color the cell in light_grey_10
                cell.paragraphs[0].runs[0].font.bold = True

            # write the first column with the name of the areas of interest
            for idx, aoi in enumerate(aois):
                cell = table.columns[0].cells[idx+1]
                cell.text = aoi

            # write the second column with the mean of dwell times of each AOI
            for idx, dwell_times in enumerate(dwell_times_df.loc['Mean'].to_numpy()):
                cell = table.columns[1].cells[idx+1]
                cell.text = str(round(dwell_times, 4))

            # write the third column with the mean of revisits of each AOI
            for idx, revisits in enumerate(revisits_df.loc['Mean'].to_numpy()):
                cell = table.columns[2].cells[idx+1]
                cell.text = str(round(revisits, 4))

            # set the vertical and horizontal alignment of all cells
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # TODO: do we need these lines
                    # cell.paragraphs[0].style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # cell.paragraphs[0].style.name = 'Table'
        '''
