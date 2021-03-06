from docx.document import Document
from bs4 import BeautifulSoup
from typing import List, Dict, Union
import pandas as pd
from docx.shared import Cm

from docx_package.picture import Picture
from docx_package.results import ResultsChapter
from docx_package.dropdown_lists import DropDownLists
from eye_tracking_package.eye_tracking import EyeTracking
from eye_tracking_package.plot import Plot


class Transitions:
    """
    Class that represents the 'Transitions' chapter and the visualization of its results.
    """

    # name of tables as they appear in the tables list
    DECISION_TABLE = 'Transitions decision table'

    # information about the headings of this chapter
    TITLE = 'Transitions'
    TITLE_STYLE = 'Heading 2'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # path to plot image files
    PARTICIPANT_FIGURE_PATH = 'Outputs/Transitions_participant{}.png'
    HEAT_MAP_FIGURE_PATH = 'Outputs/Transitions_heat_map.png'

    # caption of the pie plot figure
    CAPTION = 'Amount of transitions from an area of interest to another.'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]],
                 list_of_dataframes: List[pd.DataFrame]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
            list_of_dataframes: List of data frames containing the cGOM data of each participant
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters = parameters_dictionary
        self.cGOM_dataframes = list_of_dataframes

    def makes_plot(self):
        """
        Create heat maps to visualize the transitions percentage between the AOIs.

        One heat map for each participant is created.
        One heat map with the data of all participants are created.
        """

        # main data frame that will contain the data from the data frames from all participants
        all_transitions = pd.DataFrame()

        # create a data frame with the number of transitions for each participant and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = EyeTracking.areas_of_interest(dataframe)
            participant_transitions = EyeTracking.transitions(aois, dataframe)
            all_transitions = all_transitions.append(participant_transitions)

            # calculate the ratios and create a heat map that shows the transition percentage
            transitions_number = participant_transitions.to_numpy().sum()
            participant_transitions = participant_transitions.div(transitions_number)
            Plot.make_heatmap(data_frame=participant_transitions,
                              figure_save_path=self.PARTICIPANT_FIGURE_PATH.format(str(idx + 1)),
                              title='Transitions: participant {}'.format(str(idx + 1)),
                              xlabel='AOI destination (to)',
                              ylabel='AOI source (from)'
                              )

        # create a data frame with the total amount of transitions for each AOI and append it
        # to a data frame that will contain all transitions from all participants
        all_aois = all_transitions.columns.tolist()
        transitions_stat = pd.DataFrame()
        for idx, aoi in enumerate(all_aois):
            data_of_aoi = all_transitions[all_transitions.index == aoi]
            transitions_sum = data_of_aoi.sum().to_numpy()
            transitions_mean_df = pd.DataFrame(index=[aoi],
                                               columns=all_aois,
                                               data=[transitions_sum]
                                               )
            transitions_stat = transitions_stat.append(transitions_mean_df)

        # calculate the ratios
        transitions_number = transitions_stat.to_numpy().sum()
        transitions_stat = transitions_stat.div(transitions_number)

        # create a heat map that shows the transition percentage or do nothing if no cGOM data is provided
        if not transitions_stat.empty:
            Plot.make_heatmap(data_frame=transitions_stat,
                              figure_save_path=self.HEAT_MAP_FIGURE_PATH,
                              title='Transitions',
                              xlabel='AOI destination (to)',
                              ylabel='AOI source (from)'
                              )

    def write_chapter(self):
        """
        Write the whole chapter 'Transitions', including plot.
        """

        decision_table_index = self.tables.index(self.DECISION_TABLE)
        decision = DropDownLists.get_from_table(self.text_input_soup, decision_table_index)

        if decision[0] == 'Yes':
            self.makes_plot()

            transitions = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                         self.tables, self.picture_paths, self.parameters)

            self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)

            try:
                Picture.add_picture_and_caption(self.report,
                                                [self.HEAT_MAP_FIGURE_PATH],
                                                self.HEAT_MAP_FIGURE_PATH,
                                                self.CAPTION,
                                                width=Cm(12)
                                                )
            except FileNotFoundError:      # do nothing if no cGOM data is provided
                pass

            self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
            transitions.write_chapter()
