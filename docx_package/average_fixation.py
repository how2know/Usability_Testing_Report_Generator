from docx.document import Document
from docx.shared import Cm
from bs4 import BeautifulSoup
from typing import List, Dict, Union
import pandas as pd

from docx_package.results import ResultsChapter
from docx_package.picture import Picture
from docx_package.dropdown_lists import DropDownLists
from eye_tracking_package.eye_tracking import EyeTracking
from eye_tracking_package.plot import Plot


class AverageFixation:
    """
    Class that represents the 'Average fixation' chapter and the visualization of its results.
    """

    # name of table as it appears in the tables list
    PLOT_TYPE_TABLE = 'Average fixation plot type table'
    DECISION_TABLE = 'Average fixation decision table'

    # information about the headings of this chapter
    TITLE = 'Average fixation'
    TITLE_STYLE = 'Heading 2'
    DISCUSSION_TITLE = 'Discussion'
    DISCUSSION_STYLE = 'Heading 3'

    # path to plot image files
    PARTICIPANT_FIGURE_PATH = 'Outputs/Average_fixation_participant{}.png'
    BAR_PLOT_FIGURE_PATH = 'Outputs/Average_fixation_bar_plot.png'
    BOX_PLOT_FIGURE_PATH = 'Outputs/Average_fixation_box_plot.png'

    # caption of the pie plot figure
    BAR_PLOT_CAPTION = 'Bar plot showing the mean of the fixation duration and the 95% confidence interval.'
    BOX_PLOT_CAPTION = 'Box plot showing the mean, the 25% and 75% quartiles, and the distribution of the fixation duration.'

    def __init__(self,
                 report_document: Document,
                 text_input_document: Document,
                 text_input_soup: BeautifulSoup,
                 list_of_tables: List[str],
                 picture_paths_list: List[str],
                 parameters_dictionary: Dict[str, Union[str, int]],
                 list_of_dataframes: List[pd.DataFrame]
                 ):
        """
        Args:
            report_document: .docx file where the report is written.
            text_input_document: .docx file where all inputs are written.
            text_input_soup: BeautifulSoup of the xml of the input .docx file.
            list_of_tables: List of all table names.
            picture_paths_list: List of the path of all input pictures.
            parameters_dictionary: Dictionary of all input parameters (key = parameter name, value = parameter value)
            list_of_dataframes: List of data frames containing the cGOM data of each participant
        """

        self.report = report_document
        self.text_input = text_input_document
        self.text_input_soup = text_input_soup
        self.tables = list_of_tables
        self.picture_paths = picture_paths_list
        self.parameters = parameters_dictionary
        self.cGOM_dataframes = list_of_dataframes

    @ property
    def plot_type(self) -> str:
        """
        Returns:
            Dropdown list value of the parameter table corresponding to the plot type,
            i.e. 'Bar plot' or 'Box plot'.
        """

        plot_type_list = DropDownLists.get_from_table(self.text_input_soup,
                                                      self.tables.index(self.PLOT_TYPE_TABLE)
                                                      )
        return plot_type_list[0]

    def make_plots(self):
        """
        Create plots to visualize the fixation durations.

        One box plot for each participant is created.
        One bar plot showing a confidence interval of 95% and one box plot with the data of
        all participants are created.
        """

        # main fixation times data frame
        average_fixation_df = pd.DataFrame()

        # create a data frame with the fixation times for each participant, create a box plot with it,
        # and append it to the main data frame
        for idx, dataframe in enumerate(self.cGOM_dataframes):
            aois = EyeTracking.areas_of_interest(dataframe)
            participant_fixations = EyeTracking.fixations(aois, dataframe)

            Plot.make_boxplot(data_frame=participant_fixations,
                              figure_save_path=self.PARTICIPANT_FIGURE_PATH.format(idx + 1),
                              title='Average fixation duration: participant {}'.format(idx + 1),
                              ylabel='Fixation duration [s]',
                              xlabel='Area of interest'
                              )

            average_fixation_df = average_fixation_df.append(participant_fixations, ignore_index=True)

        # create a bar plot and a box plot with the fixations of all participants or
        # do nothing if no cGOM data is provided
        try:
            Plot.make_boxplot(data_frame=average_fixation_df,
                              figure_save_path=self.BOX_PLOT_FIGURE_PATH,
                              title='Average fixation duration',
                              ylabel='Fixation duration [s]',
                              xlabel='Area of interest'
                              )
            Plot.make_barplot(data_frame=average_fixation_df,
                              figure_save_path=self.BAR_PLOT_FIGURE_PATH,
                              title='Average fixation duration',
                              ylabel='Fixation duration [s]',
                              xlabel='Area of interest'
                              )
        except ValueError:
            pass

    def write_chapter(self):
        """
        Write the whole chapter 'Average fixation', including the chosen plot.
        """

        decision_table_index = self.tables.index(self.DECISION_TABLE)
        decision = DropDownLists.get_from_table(self.text_input_soup, decision_table_index)

        if decision[0] == 'Yes':
            self.make_plots()

            time_on_tasks = ResultsChapter(self.report, self.text_input, self.text_input_soup, self.TITLE,
                                           self.tables, self.picture_paths, self.parameters)

            self.report.add_paragraph(self.TITLE, self.TITLE_STYLE)

            # add bar plot or box plot depending on the choice of plot type or do nothing if no cGOM data is provided
            try:
                if self.plot_type == 'Bar plot':
                    Picture.add_picture_and_caption(self.report,
                                                    [self.BAR_PLOT_FIGURE_PATH],
                                                    self.BAR_PLOT_FIGURE_PATH,
                                                    self.BAR_PLOT_CAPTION,
                                                    width=Cm(12)
                                                    )
                if self.plot_type == 'Box plot':
                    Picture.add_picture_and_caption(self.report,
                                                    [self.BOX_PLOT_FIGURE_PATH],
                                                    self.BOX_PLOT_FIGURE_PATH,
                                                    self.BOX_PLOT_CAPTION,
                                                    width=Cm(12)
                                                    )
            except FileNotFoundError:
                pass

            self.report.add_paragraph(self.DISCUSSION_TITLE, self.DISCUSSION_STYLE)
            time_on_tasks.write_chapter()
