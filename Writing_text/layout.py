
#####     LAYOUT DEFINITION     #####

# Some functions that define the layout and formatting of the report are implemented in this module.
# It includes: document layout, header, footer, styles, tables, ...

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement

# define some colors
black = RGBColor(0, 0, 0)                  # Hex: 000000
black_35 = RGBColor(90, 90, 90)            # Hex: 5A5A5A
light_grey_10 = RGBColor(208, 206, 206)    # Hex: D0CECE

# define the page setup of as section as default A4 setup
def define_page_format(section):
    section.orientation = WD_ORIENT.PORTRAIT   # orientation of the page: portrait
    section.page_width = Cm(21)                # page width: 21 cm
    section.page_height = Cm(29.7)             # page height: 29,7 cm
    section.top_margin = Cm(2.5)               # top margin: 2,5 cm
    section.bottom_margin = Cm(2.5)            # bottom margin: 2,5 cm
    section.right_margin = Cm(2.5)             # right margin: 2,5 cm
    section.left_margin = Cm(2.5)              # left margin: 2,5 cm

# define the characteristics of a style of a document
def style_definition(document, name, font, size, color, alignment, italic, bold):
    style = document.styles[name]                                   # name of the style you want to define
    try:
        style.element.xpath('w:rPr/w:rFonts')[0].attrib.clear()     # this let us modify the font of some styles that are kind of "blocked"
    except IndexError:                                              # pass the styles for which this does not work
        pass
    style.font.name = font                                          # font name
    style.font.size = Pt(size)                                      # font size (in pt)
    style.font.color.rgb = color                                    # font color
    style.paragraph_format.alignment = alignment                    # alignment of the paragraph (left, right, center, justify)
    style.font.italic = italic                                      # boolean to know if it should be written in italic
    style.font.bold = bold                                          # boolean to know if it should be written in bold

# define all styles used in a document using the function style_definition
def define_all_styles(document):
    style_definition(document, 'Title', 'Calibri Light', 32, black, WD_ALIGN_PARAGRAPH.CENTER, False, False)
    style_definition(document, 'Subtitle', 'Calibri Light', 24, black_35, WD_ALIGN_PARAGRAPH.CENTER, False, False)
    style_definition(document, 'Heading 1', 'Calibri', 16, black, WD_ALIGN_PARAGRAPH.LEFT, False, True)
    style_definition(document, 'Heading 2', 'Calibri Light', 14, black, WD_ALIGN_PARAGRAPH.LEFT, False, True)
    style_definition(document, 'Normal', 'Calibri', 11, black, WD_ALIGN_PARAGRAPH.JUSTIFY, False, False)

    '''Créer un styler pour les tables, headers, ... et tout le reste. Chaque truc différent doit avoir un fucking style men'''

# create a header with two lines for a section
def create_header(section, first_line_content, second_line_content):
    header = section.header
    header.is_linked_to_previous = False
    header_first_line = header.paragraphs[0]
    header_first_line.paragraph_format.tab_stops.clear_all()
    header_first_line.paragraph_format.tab_stops.add_tab_stop(Cm(0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
    header_first_line.paragraph_format.tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES)
    header_first_line.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    header_first_line.text = first_line_content
    header_second_line = header.add_paragraph()
    header_second_line.paragraph_format.tab_stops.add_tab_stop(Cm(0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
    header_second_line.paragraph_format.tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES)
    header_second_line.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    header_second_line.text = second_line_content

# techniques to make something bold in a table
# copied from: https://stackoverflow.com/questions/37757203/making-cells-bold-in-a-table-using-python-docx
'''
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
'''

# define table style
def define_table_style(table):
    table.style = 'Table Grid'                            # set the table style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER           # set the table alignment

    # set the vertical and horizontal alignment of all cells
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # set the shading of the first row to light_grey_10 (RGB Hex: D0CECE)
    # copied from: https://groups.google.com/forum/#!topic/python-docx/-c3OrRHA3qo
    for cell in table.rows[0].cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="D0CECE"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)



# insert an horizontal border under a given paragraph
# copied from: https://github.com/python-openxml/python-docx/issues/105
def insert_horizontal_border(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)